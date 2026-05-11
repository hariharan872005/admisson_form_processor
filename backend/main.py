from fastapi import FastAPI, UploadFile, File, Depends, HTTPException
from fastapi.middleware.cors import CORSMiddleware
import pandas as pd
import io
import re
from sqlalchemy.orm import Session
from typing import List, Optional

import models, schemas, database
from database import engine

# Create tables if they don't exist (never drop)
models.Base.metadata.create_all(bind=engine)

app = FastAPI(title="Admission Form Processor API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── CSV column name aliases → schema field names ───────────────────────────────
CSV_COLUMN_MAP = {
    # Personal
    "name": "name", "full name": "name", "student name": "name",
    "dob": "date_of_birth", "date of birth": "date_of_birth", "birth date": "date_of_birth",
    "gender": "gender", "sex": "gender",
    "blood group": "blood_group", "bloodgroup": "blood_group",
    "aadhaar": "aadhaar_number", "aadhaar number": "aadhaar_number", "aadhar": "aadhaar_number",
    "nationality": "nationality",
    "religion": "religion",
    "community": "community",
    "caste": "caste",
    # Contact
    "mobile": "mobile_number", "mobile number": "mobile_number", "phone": "mobile_number",
    "email": "email", "email id": "email", "email address": "email",
    "address": "residential_address", "residential address": "residential_address",
    "parent contact": "parent_contact_number", "parent mobile": "parent_contact_number",
    "parent's contact number": "parent_contact_number",
    # Parent / Guardian
    "father name": "father_name", "father's name": "father_name",
    "father occupation": "father_occupation", "father's occupation": "father_occupation",
    "mother name": "mother_name", "mother's name": "mother_name",
    "mother occupation": "mother_occupation", "mother's occupation": "mother_occupation",
    "annual income": "annual_income", "income": "annual_income",
    "guardian name": "guardian_name", "guardian": "guardian_name",
    "guardian relation": "guardian_relation",
    # Academic
    "regno": "registration_number", "reg no": "registration_number",
    "registration number": "registration_number", "roll number": "registration_number",
    "previous school": "previous_school", "school name": "previous_school",
    "board": "board_university", "university": "board_university",
    "board/university": "board_university",
    "marks": "marks_percentage", "percentage": "marks_percentage",
    "marks / percentage / cgpa": "marks_percentage", "cgpa": "marks_percentage",
    "year of passing": "year_of_passing", "passing year": "year_of_passing",
    "tc number": "tc_number", "transfer certificate": "tc_number",
    # Course
    "course": "course_department", "department": "course_department",
    "course / department": "course_department",
    "medium": "medium_of_instruction", "medium of instruction": "medium_of_instruction",
    "admission category": "admission_category", "category": "admission_category",
    # Bank
    "account number": "bank_account_number", "bank account": "bank_account_number",
    "ifsc": "ifsc_code", "ifsc code": "ifsc_code",
    "bank name": "bank_name", "bank": "bank_name",
    # Other
    "hostel": "hostel_required", "hostel requirement": "hostel_required",
    "transport": "transport_required", "transport requirement": "transport_required",
    "scholarship": "scholarship_details", "scholarship details": "scholarship_details",
    "emergency contact": "emergency_contact_name",
    "emergency contact name": "emergency_contact_name",
    "emergency contact number": "emergency_contact_number",
    "emergency mobile": "emergency_contact_number",
}

# ── Regex patterns for PDF/DOCX extraction ────────────────────────────────────
FIELD_PATTERNS = {
    "name": re.compile(
        r'(?:full\s*name|student\s*name|name)\s*[:\-]?\s*([A-Za-z][A-Za-z\s.\'-]{1,40}?)(?=\s*\n|\s{2,}|\s*(?:date|dob|gender|mobile|phone|email|reg|roll|father|mother|blood|aadhaar|address|course|bank|hostel|transport|$))',
        re.IGNORECASE),
    "date_of_birth": re.compile(
        r'(?:date\s*of\s*birth|dob|birth\s*date)\s*[:\-]?\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})',
        re.IGNORECASE),
    "gender": re.compile(
        r'(?:gender|sex)\s*[:\-]?\s*(male|female|other|transgender)',
        re.IGNORECASE),
    "blood_group": re.compile(
        r'(?:blood\s*group)\s*[:\-]?\s*([ABO]{1,2}[+-])',
        re.IGNORECASE),
    "aadhaar_number": re.compile(
        r'(?:aadhaar|aadhar)\s*(?:number|no|#)?\s*[:\-]?\s*(\d{4}\s?\d{4}\s?\d{4})',
        re.IGNORECASE),
    "nationality": re.compile(
        r'(?:nationality)\s*[:\-]?\s*([A-Za-z]{3,30})',
        re.IGNORECASE),
    "religion": re.compile(
        r'(?:religion)\s*[:\-]?\s*([A-Za-z]{3,30})',
        re.IGNORECASE),
    "community": re.compile(
        r'(?:community)\s*[:\-]?\s*([A-Za-z\s]{2,30})',
        re.IGNORECASE),
    "caste": re.compile(
        r'(?:caste)\s*[:\-]?\s*([A-Za-z\s]{2,30})',
        re.IGNORECASE),
    "mobile_number": re.compile(
        r'(?:mobile|phone|contact)\s*(?:number|no|#)?\s*[:\-]?\s*(\+?[\d\s\-]{10,15})',
        re.IGNORECASE),
    "email": re.compile(
        r'[\w.\-+]+@[\w.\-]+\.[a-zA-Z]{2,}'),
    "residential_address": re.compile(
        r'(?:residential\s*address|address)\s*[:\-]?\s*(.{10,150})',
        re.IGNORECASE),
    "parent_contact_number": re.compile(
        r'(?:parent(?:\'?s)?\s*(?:contact|mobile|phone))\s*(?:number|no)?\s*[:\-]?\s*(\+?[\d\s\-]{10,15})',
        re.IGNORECASE),
    "father_name": re.compile(
        r'(?:father(?:\'?s)?\s*name)\s*[:\-]?\s*([A-Za-z][A-Za-z\s.\'-]{1,50})',
        re.IGNORECASE),
    "father_occupation": re.compile(
        r'(?:father(?:\'?s)?\s*occupation)\s*[:\-]?\s*([A-Za-z\s]{2,50})',
        re.IGNORECASE),
    "mother_name": re.compile(
        r'(?:mother(?:\'?s)?\s*name)\s*[:\-]?\s*([A-Za-z][A-Za-z\s.\'-]{1,50})',
        re.IGNORECASE),
    "mother_occupation": re.compile(
        r'(?:mother(?:\'?s)?\s*occupation)\s*[:\-]?\s*([A-Za-z\s]{2,50})',
        re.IGNORECASE),
    "annual_income": re.compile(
        r'(?:annual\s*income|income)\s*[:\-]?\s*([\d,\.]+(?:\s*(?:lakhs?|lpa|rs\.?|₹))?)',
        re.IGNORECASE),
    "guardian_name": re.compile(
        r'(?:guardian(?:\'?s)?\s*name|guardian)\s*[:\-]?\s*([A-Za-z][A-Za-z\s.\'-]{1,50})',
        re.IGNORECASE),
    "registration_number": re.compile(
        r'(?:reg(?:istration)?\s*(?:no|number|#)?|roll\s*(?:no|number))\s*[:\-]?\s*([A-Z0-9\-/]{3,20})',
        re.IGNORECASE),
    "previous_school": re.compile(
        r'(?:previous\s*(?:school|college)|school\s*name|college\s*name)\s*[:\-]?\s*(.{3,80})',
        re.IGNORECASE),
    "board_university": re.compile(
        r'(?:board|university)\s*[:\-]?\s*([A-Za-z\s]{3,60})',
        re.IGNORECASE),
    "marks_percentage": re.compile(
        r'(?:marks|percentage|cgpa|%)\s*[:\-]?\s*([\d.]+\s*%?)',
        re.IGNORECASE),
    "year_of_passing": re.compile(
        r'(?:year\s*of\s*passing|passing\s*year)\s*[:\-]?\s*(\d{4})',
        re.IGNORECASE),
    "tc_number": re.compile(
        r'(?:tc\s*(?:no|number)|transfer\s*certificate)\s*[:\-]?\s*([A-Z0-9\-/]{2,20})',
        re.IGNORECASE),
    "course_department": re.compile(
        r'(?:course|department|programme|program)\s*[:\-]?\s*([A-Za-z\s.&]{2,60})',
        re.IGNORECASE),
    "medium_of_instruction": re.compile(
        r'(?:medium\s*(?:of\s*instruction)?)\s*[:\-]?\s*([A-Za-z]{3,30})',
        re.IGNORECASE),
    "admission_category": re.compile(
        r'(?:admission\s*category|category)\s*[:\-]?\s*([A-Za-z\s]{2,40})',
        re.IGNORECASE),
    "bank_account_number": re.compile(
        r'(?:account\s*(?:no|number)|bank\s*account)\s*[:\-]?\s*(\d{9,18})',
        re.IGNORECASE),
    "ifsc_code": re.compile(
        r'(?:ifsc\s*(?:code)?)\s*[:\-]?\s*([A-Z]{4}0[A-Z0-9]{6})',
        re.IGNORECASE),
    "bank_name": re.compile(
        r'(?:bank\s*name|bank)\s*[:\-]?\s*([A-Za-z\s]{3,50})',
        re.IGNORECASE),
    "hostel_required": re.compile(
        r'(?:hostel\s*(?:requirement|required|needed)?)\s*[:\-]?\s*(yes|no)',
        re.IGNORECASE),
    "transport_required": re.compile(
        r'(?:transport\s*(?:requirement|required|needed)?)\s*[:\-]?\s*(yes|no)',
        re.IGNORECASE),
    "scholarship_details": re.compile(
        r'(?:scholarship\s*(?:details)?)\s*[:\-]?\s*(.{2,100})',
        re.IGNORECASE),
    "emergency_contact_name": re.compile(
        r'(?:emergency\s*contact\s*(?:name)?)\s*[:\-]?\s*([A-Za-z][A-Za-z\s.\'-]{1,50})',
        re.IGNORECASE),
    "emergency_contact_number": re.compile(
        r'(?:emergency\s*(?:contact\s*)?(?:number|mobile|phone))\s*[:\-]?\s*(\+?[\d\s\-]{10,15})',
        re.IGNORECASE),
}


# ── Text extraction helpers ───────────────────────────────────────────────────

def extract_text_from_pdf(contents: bytes) -> str:
    try:
        import fitz
        doc = fitz.open(stream=contents, filetype="pdf")
        return "\n".join(page.get_text() for page in doc)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading PDF: {str(e)}")


def extract_text_from_docx(contents: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(contents))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading DOCX: {str(e)}")


def parse_students_from_text(text: str):
    """
    Extract student fields from PDF text.
    Handles two layouts:
      1. Table-style: 'Field\nValue\nField\nValue...' (PyMuPDF output from table PDFs)
      2. Inline-style: 'Field: Value\nField: Value...'
    """
    extracted_data = []
    errors = []

    # ── Known field label → schema key mapping ────────────────────────────────
    LABEL_MAP = {
        "full name": "name", "student name": "name", "name": "name",
        "date of birth": "date_of_birth", "dob": "date_of_birth",
        "gender": "gender", "sex": "gender",
        "blood group": "blood_group",
        "aadhaar number": "aadhaar_number", "aadhaar": "aadhaar_number",
        "aadhar number": "aadhaar_number", "aadhar": "aadhaar_number",
        "nationality": "nationality",
        "religion": "religion",
        "community": "community",
        "caste": "caste",
        "mobile number": "mobile_number", "mobile": "mobile_number",
        "phone": "mobile_number", "contact number": "mobile_number",
        "email id": "email", "email": "email", "email address": "email",
        "residential address": "residential_address", "address": "residential_address",
        "parent's contact number": "parent_contact_number",
        "parent contact number": "parent_contact_number",
        "parent mobile": "parent_contact_number",
        "father's name": "father_name", "father name": "father_name",
        "father's occupation": "father_occupation", "father occupation": "father_occupation",
        "mother's name": "mother_name", "mother name": "mother_name",
        "mother's occupation": "mother_occupation", "mother occupation": "mother_occupation",
        "annual income": "annual_income", "income": "annual_income",
        "guardian name": "guardian_name", "guardian's name": "guardian_name",
        "guardian relation": "guardian_relation", "guardian relationship": "guardian_relation",
        "registration number": "registration_number", "reg no": "registration_number",
        "regno": "registration_number", "roll number": "registration_number",
        "previous school / college": "previous_school",
        "previous school": "previous_school", "previous college": "previous_school",
        "school name": "previous_school", "college name": "previous_school",
        "board / university": "board_university",
        "board": "board_university", "university": "board_university",
        "marks / percentage / cgpa": "marks_percentage",
        "marks": "marks_percentage", "percentage": "marks_percentage", "cgpa": "marks_percentage",
        "year of passing": "year_of_passing", "passing year": "year_of_passing",
        "tc number": "tc_number", "transfer certificate": "tc_number",
        "course / department": "course_department",
        "course": "course_department", "department": "course_department",
        "medium of instruction": "medium_of_instruction", "medium": "medium_of_instruction",
        "admission category": "admission_category", "category": "admission_category",
        "bank account number": "bank_account_number", "account number": "bank_account_number",
        "ifsc code": "ifsc_code", "ifsc": "ifsc_code",
        "bank name": "bank_name", "bank": "bank_name",
        "hostel required": "hostel_required", "hostel": "hostel_required",
        "transport required": "transport_required", "transport": "transport_required",
        "scholarship details": "scholarship_details", "scholarship": "scholarship_details",
        "emergency contact name": "emergency_contact_name",
        "emergency contact": "emergency_contact_name",
        "emergency contact number": "emergency_contact_number",
        "emergency mobile": "emergency_contact_number",
    }

    # Build a set of all known labels for fast lookup
    known_labels = set(LABEL_MAP.keys())

    lines = [l.strip() for l in text.splitlines() if l.strip()]

    # Skip header lines like "Field  Value" or "Modified Student Details..."
    skip_patterns = re.compile(
        r'^(field|value|modified student|student details|education details|unchanged)$',
        re.IGNORECASE
    )

    record = {}

    # ── Strategy 1: Table layout — alternating label/value lines ─────────────
    # Detect if most odd lines are known labels
    candidate_labels = [l for l in lines if l.lower().strip(":.") in known_labels]
    if len(candidate_labels) >= 5:
        # Table-style: lines alternate between label and value
        i = 0
        while i < len(lines):
            line = lines[i]
            label_key = line.lower().strip(":. ")
            if label_key in LABEL_MAP:
                field = LABEL_MAP[label_key]
                # Value is the next non-label line
                if i + 1 < len(lines):
                    next_line = lines[i + 1]
                    next_key = next_line.lower().strip(":. ")
                    # Only use as value if next line is NOT itself a label
                    if next_key not in known_labels and not skip_patterns.match(next_line):
                        record[field] = next_line.strip()
                        i += 2
                        continue
            i += 1
    else:
        # ── Strategy 2: Inline layout — "Label: Value" on same line ──────────
        for line in lines:
            if skip_patterns.match(line):
                continue
            # Try "Label: Value" or "Label - Value"
            m = re.match(r'^(.+?)\s*[:\-]\s*(.+)$', line)
            if m:
                label_key = m.group(1).strip().lower()
                value = m.group(2).strip()
                if label_key in LABEL_MAP and value:
                    record[LABEL_MAP[label_key]] = value

    if not record:
        return [], [{"block": 1, "error": "No fields could be extracted from the document.",
                     "raw": text[:300]}]

    # If we have a valid record, validate and return
    try:
        student = schemas.StudentData(**record)
        extracted_data.append(
            student.dict() if hasattr(student, "dict") else student.model_dump()
        )
    except Exception as e:
        errors.append({"block": 1, "error": str(e), "raw": str(record)[:300]})

    return extracted_data, errors


# ── DataFrame processing ──────────────────────────────────────────────────────

def _process_dataframe(df: pd.DataFrame):
    """Map CSV/Excel columns to schema fields and validate each row."""
    extracted_data = []
    errors = []

    # Normalise column names
    col_map = {}
    for col in df.columns:
        normalised = col.strip().lower()
        if normalised in CSV_COLUMN_MAP:
            col_map[col] = CSV_COLUMN_MAP[normalised]

    if not col_map:
        return {"status": "error",
                "message": "No recognisable columns found. Check the Help page for expected column names."}

    df = df.fillna("")

    for index, row in df.iterrows():
        record = {}
        for orig_col, field_name in col_map.items():
            val = str(row.get(orig_col, "")).strip()
            if val:
                record[field_name] = val

        if not record.get("name"):
            errors.append({"row": index + 1, "error": "Missing required field: name"})
            continue

        try:
            student = schemas.StudentData(**record)
            extracted_data.append(
                student.dict() if hasattr(student, "dict") else student.model_dump()
            )
        except Exception as e:
            errors.append({"row": index + 1, "error": str(e)})

    return {"status": "success", "valid_data": extracted_data, "errors": errors,
            "source_type": "structured"}


# ── Routes ────────────────────────────────────────────────────────────────────

@app.post("/process-form/")
async def process_form(file: UploadFile = File(...)):
    contents = await file.read()
    filename = file.filename.lower()

    if filename.endswith(".csv"):
        try:
            df = pd.read_csv(io.StringIO(contents.decode("utf-8")))
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error reading CSV: {str(e)}")
        return _process_dataframe(df)

    elif filename.endswith((".xlsx", ".xls")):
        try:
            df = pd.read_excel(io.BytesIO(contents))
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error reading Excel: {str(e)}")
        return _process_dataframe(df)

    elif filename.endswith(".pdf"):
        text = extract_text_from_pdf(contents)
        extracted_data, errors = parse_students_from_text(text)
        return {"status": "success", "valid_data": extracted_data, "errors": errors,
                "source_type": "pdf"}

    elif filename.endswith((".docx", ".doc")):
        text = extract_text_from_docx(contents)
        extracted_data, errors = parse_students_from_text(text)
        return {"status": "success", "valid_data": extracted_data, "errors": errors,
                "source_type": "docx"}

    else:
        raise HTTPException(status_code=400,
                            detail="Unsupported format. Use PDF, DOCX, CSV, or Excel.")


@app.post("/export/")
def export_data(students: List[schemas.StudentData], db: Session = Depends(database.get_db)):
    exported_count = 0

    for s in students:
        data = s.dict() if hasattr(s, "dict") else s.model_dump()
        db_student = models.Student(**data)
        db.add(db_student)
        exported_count += 1

    db.commit()
    return {"status": "success", "exported": exported_count, "skipped_duplicates": 0}


@app.get("/students")
def get_students(db: Session = Depends(database.get_db)):
    students = db.query(models.Student).all()
    result = []
    for s in students:
        row = {c.name: getattr(s, c.name) for c in models.Student.__table__.columns}
        result.append(row)
    return result


@app.delete("/students/{student_id}")
def delete_student(student_id: int, db: Session = Depends(database.get_db)):
    student = db.query(models.Student).filter(models.Student.id == student_id).first()
    if not student:
        raise HTTPException(status_code=404, detail="Student not found.")
    db.delete(student)
    db.commit()
    return {"status": "success", "message": f"Deleted student id {student_id}"}
